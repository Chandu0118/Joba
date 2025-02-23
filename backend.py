from docx import Document
import requests
import os
from dotenv import load_dotenv
from grok_api import GrokModel

# Load environment variables
load_dotenv()
grok_api_key = os.getenv("GROK_API_KEY")

def process_cv_and_generate_cover_letter(cv_path, job_input, selected_llm, output_language):
    # Load CV
    cv_doc = Document(cv_path)
    cv_text = "\n".join([para.text for para in cv_doc.paragraphs])

    # Process Job Description
    if job_input.startswith("http"):
        response = requests.get(job_input)
        job_description = response.text
    else:
        job_description = job_input

    # Initialize Grok only
    if selected_llm != "Grok":
        raise ValueError("Only Grok LLM is supported for now.")
    llm = GrokModel(model_name="grok-2", api_key=grok_api_key)  # Using Grok 2 as specified

    # Generate CV Updates
    prompt_cv = (
        f"Update the following CV based on the job description. "
        f"Do not carry over information from the job description unless explicitly mentioned. "
        f"Output language: {output_language}.\n\n"
        f"CV:\n{cv_text}\n\nJob Description:\n{job_description}"
    )
    updated_cv_text = llm(prompt_cv)

    # Save Updated CV
    updated_cv_doc = Document()
    updated_cv_doc.add_paragraph(updated_cv_text)
    if not os.path.exists("output"):
        os.makedirs("output")
    updated_cv_path = os.path.join("output", "updated_cv.docx")
    updated_cv_doc.save(updated_cv_path)

    # Generate Cover Letter
    prompt_cover = (
        f"Write a professional cover letter in {output_language} for the following job description:\n\n"
        f"{job_description}\n\nCV:\n{cv_text}"
    )
    cover_letter = llm(prompt_cover)

    # Save Cover Letter as .docx
    cover_letter_doc = Document()
    cover_letter_doc.add_paragraph(cover_letter)
    cover_letter_path = os.path.join("output", "cover_letter.docx")
    cover_letter_doc.save(cover_letter_path)

    return updated_cv_path, cover_letter_path

def analyze_feedback(original_cv_path, updated_cv_path):
    original_doc = Document(original_cv_path)
    updated_doc = Document(updated_cv_path)

    original_text = "\n".join([para.text for para in original_doc.paragraphs])
    updated_text = "\n".join([para.text for para in updated_doc.paragraphs])

    # Compare texts and log differences for learning
    differences = set(updated_text.split()) - set(original_text.split())
    with open("feedback_log.txt", "a", encoding="utf-8") as log:
        log.write(f"Differences in CV update: {', '.join(differences)}\n")
    return ", ".join(differences)
